"""Generate the jasper-deploy LinkedIn post as a branded DOCX.

ASCII-only content (no em dashes, curly quotes, or arrows) per house style.
Actian brand colors: Software Blue #000032, Tech Blue #3C91FF, Software Teal #2EC0CB.
"""
import os
import subprocess
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Inches

SOFTWARE_BLUE = RGBColor(0x00, 0x00, 0x32)
TECH_BLUE = RGBColor(0x3C, 0x91, 0xFF)
SOFTWARE_TEAL = RGBColor(0x2E, 0xC0, 0xCB)
GREY = RGBColor(0x55, 0x55, 0x55)

OUT = os.path.join(os.path.expanduser("~"), "Downloads",
                   "LinkedIn_Post_jasper_deploy_Plugin.docx")


def style_base(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = SOFTWARE_BLUE
    pf = normal.paragraph_format
    pf.space_after = Pt(8)
    pf.line_spacing = 1.15


def heading(doc, text, size=16, color=SOFTWARE_BLUE, space_before=14):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = color
    return p


def rule(doc, color=SOFTWARE_TEAL):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("_" * 74)
    r.font.size = Pt(8)
    r.font.color.rgb = color
    return p


def body(doc, text, bold=False, italic=False, color=None, size=11,
         align=None, space_after=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    r.font.color.rgb = color or SOFTWARE_BLUE
    return p


def bullet(doc, text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
        r.font.color.rgb = SOFTWARE_BLUE
    r = p.add_run(text)
    r.font.color.rgb = SOFTWARE_BLUE
    return p


def main():
    doc = Document()
    for s in doc.sections:
        s.left_margin = s.right_margin = Inches(1.0)
        s.top_margin = s.bottom_margin = Inches(0.9)
    style_base(doc)

    # ---- Title block -------------------------------------------------
    body(doc, "LINKEDIN POST", bold=True, size=9, color=TECH_BLUE,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    body(doc, "Building and Shipping the jasper-deploy Plugin",
         bold=True, size=20, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    body(doc,
         "Ten weeks of agent-built automation, held to production standards",
         italic=True, size=11, color=GREY,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    body(doc, "Robert Gorsuch  |  Prepared August 17, 2026",
         size=9, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    rule(doc)

    # ---- Version A: the main post ------------------------------------
    heading(doc, "Version A - Full post (recommended)", size=14,
            color=TECH_BLUE, space_before=0)
    body(doc, "Approx. 320 words. Paste directly into LinkedIn. Line breaks "
              "are intentional; LinkedIn collapses paragraphs without them.",
         italic=True, size=9, color=GREY, space_after=10)

    post_a = [
        "I spent ten weeks teaching an AI agent to run an enterprise "
        "reporting platform. The hard part was never the building. It was "
        "proving the work was real.",

        "The result is jasper-deploy: a Claude Code plugin that automates the "
        "full JasperReports Server lifecycle. Scaffold a report from a SQL "
        "query, lint it, compile it, deploy it over REST, run it, and verify "
        "the bytes that come back. Then dashboards, Domains, OLAP cubes, "
        "input controls, schedules, alerts, themes, users, roles, "
        "permissions, and promotion between environments.",

        "It started on June 1 as a single happy path. Everything after that "
        "grew the same way: attempt a real task against a live server, fix "
        "whatever breaks, and convert the fix into executable tooling so no "
        "future session ever relearns it.",

        "Four principles did the heavy lifting:",

        "1. Verify against the live system, never the documentation. Nothing "
        "was marked as working until it round-tripped end to end: HTTP "
        "status, PDF magic bytes, CSV row counts, rendered-pixel baselines. "
        "Import and export were proven by deleting the live resource and "
        "re-importing it.",

        "2. Every failure becomes a guardrail, not a memory. Painful "
        "discoveries were promoted in order of strength: a note, then a "
        "symptom-indexed reference, then a lint rule, then a hard gate inside "
        "the deploy script. My favorite: a hand-built dashboard posted to the "
        "API stores cleanly with a 201 and then renders completely blank. "
        "That one cost a day. It can never cost anyone a day again.",

        "3. Separate verified from doc-only. Every endpoint in the reference "
        "map is labeled by whether it was actually exercised. The server's own "
        "live WADL outranks the vendor PDF.",

        "4. Keep the agent-facing index lean. The skill file is a capability "
        "map, one line per task. Depth lives in linked references, and a CI "
        "check fails the build if a link or a script goes missing.",

        "Where it landed: 51 scripts, 29 reference docs, a 24-step smoke test, "
        "a Pester unit suite, golden-image visual baselines, a full-history "
        "secret scan in CI, Apache-2.0, and four slash commands. Now shipping "
        "as an installable plugin at v1.1.0.",

        "AI made the build fast. Verification is what made it trustworthy. If "
        "you are training skills or plugins for real systems, budget for the "
        "second part.",
    ]
    for para in post_a:
        body(doc, para, space_after=10)

    body(doc,
         "#AI #Automation #ClaudeCode #DataEngineering #BusinessIntelligence "
         "#Jaspersoft #DevOps #AIAgents",
         color=TECH_BLUE, space_after=6)

    rule(doc)

    # ---- Version B: short punchy -------------------------------------
    heading(doc, "Version B - Short and punchy", size=14, color=TECH_BLUE)
    body(doc, "Approx. 150 words. Higher completion rate; better if you want "
              "the detail to live in the comments.",
         italic=True, size=9, color=GREY, space_after=10)

    post_b = [
        "A hand-built dashboard posted to the JasperReports API returns 201 "
        "Created. Then it renders completely blank. No error, no warning, "
        "nothing in the logs.",

        "That bug cost me a day. It will never cost anyone a day again, "
        "because the fix is now a hard gate inside a deploy script.",

        "That is the whole philosophy behind jasper-deploy, the Claude Code "
        "plugin I just shipped at v1.1.0. It automates the entire "
        "JasperReports Server lifecycle: scaffold from SQL, lint, compile, "
        "deploy, verify, dashboards, Domains, OLAP, scheduling, admin, and "
        "environment promotion.",

        "Ten weeks. 51 scripts. 29 reference docs. A 24-step smoke test that "
        "asserts every stage against a live server. Unit tests and a "
        "full-history secret scan in CI.",

        "The rule that made it work: every failure becomes a guardrail, not a "
        "memory. An AI agent can rebuild your knowledge in seconds. It cannot "
        "remember your scar tissue unless you make the scar tissue "
        "executable.",

        "That is the difference between an AI demo and AI infrastructure.",
    ]
    for para in post_b:
        body(doc, para, space_after=10)

    body(doc, "#AI #AIAgents #Automation #ClaudeCode #BusinessIntelligence "
              "#PlatformEngineering",
         color=TECH_BLUE, space_after=6)

    rule(doc)

    # ---- Supporting facts --------------------------------------------
    heading(doc, "Supporting facts (for comments or follow-up posts)",
            size=14, color=TECH_BLUE)

    bullet(doc, "June 1 to August 7, 2026. Roughly 70 skill-related commits "
                "across the repository.", bold_lead="Build span: ")
    bullet(doc, "51 PowerShell and Python scripts, 29 reference documents, "
                "3 JSON Schemas, a snapshot of the server's live WADL.",
           bold_lead="Surface: ")
    bullet(doc, "24-step smoke test under a throwaway folder, plus offline "
                "prechecks that run first so a broken script never reaches "
                "the server.", bold_lead="Regression gate: ")
    bullet(doc, "A JR7 lint pass runs automatically inside the deploy script. "
                "Valid element names were extracted from the JasperReports "
                "7.0.6 source annotations, not from prose.",
           bold_lead="Pre-deploy gate: ")
    bullet(doc, "gitleaks full-history scan, doc and link consistency check, "
                "Pester unit tests on a Windows plus Ubuntu matrix.",
           bold_lead="CI: ")
    bullet(doc, "Claude Code sessions and a self-service web wizard call the "
                "identical scripts, so the UI and the agent cannot drift "
                "apart on behavior.", bold_lead="Two consumers, one core: ")
    bullet(doc, "v1.1.0 trimmed the install payload from the whole working "
                "repository to the plugin only, added four slash commands, "
                "and moved every machine-specific fact out of the skill file "
                "into an optional local overlay.",
           bold_lead="Packaging discipline: ")

    heading(doc, "War stories worth telling", size=14, color=TECH_BLUE)
    bullet(doc, "Stores with a 201, renders blank. Composition has to go "
                "through export, inject, import.",
           bold_lead="The dashboard that lies: ")
    bullet(doc, "The server's SQL security validator rejects any query that "
                "does not begin with SELECT, so a CTE fails at deploy. The "
                "linter now blocks a leading WITH and points at the "
                "nested-subquery rewrite.", bold_lead="The CTE trap: ")
    bullet(doc, "A clean local compile proves nothing. A strict JSON parser "
                "rejects unknown elements as an opaque 400 at fill time, "
                "long after compile succeeded.",
           bold_lead="The compile that means nothing: ")
    bullet(doc, "Different chart types accept different plot attributes, and "
                "one accepts none at all. Now a per-kind lint rule.",
           bold_lead="Charts are not interchangeable: ")

    heading(doc, "The transferable method", size=14, color=TECH_BLUE)
    body(doc, "Product-agnostic. This is the part worth reusing.",
         italic=True, size=9, color=GREY, space_after=8)
    for n, line in enumerate([
        "Start with one verified happy path, not a framework.",
        "Grow only by attempting real tasks against the real system.",
        "Convert every failure into the strongest guardrail you can afford: "
        "note, then reference, then lint rule, then hard gate.",
        "Keep the agent-facing index lean and push depth into linked "
        "references, then make a CI check own their consistency.",
        "Label knowledge as verified or doc-only, and prefer the system's own "
        "introspection over vendor prose.",
        "Build the regression gate in step with the surface area, never "
        "after.",
    ], start=1):
        bullet(doc, line, bold_lead="%d. " % n)

    doc.save(OUT)
    print(OUT)
    return OUT


if __name__ == "__main__":
    path = main()
    if "--open" in sys.argv:
        os.startfile(path)  # noqa: S606 (Windows shell open)
